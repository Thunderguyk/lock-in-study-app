"""
File processing utilities for Lock-In App
Handles PDF, DOCX, and TXT file text extraction
"""

import io
import re
from typing import Dict, Any
import streamlit as st

try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False

try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

def extract_text_from_pdf(file_content: bytes) -> str:
    """Extract text from PDF file using available libraries"""
    
    if PYMUPDF_AVAILABLE:
        try:
            doc = fitz.open(stream=file_content, filetype="pdf")
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except Exception as e:
            st.error(f"PyMuPDF extraction failed: {e}")
    
    if PDFPLUMBER_AVAILABLE:
        try:
            with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                text = ""
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                return text
        except Exception as e:
            st.error(f"pdfplumber extraction failed: {e}")
    
    return "PDF text extraction not available. Please install PyMuPDF or pdfplumber."

def extract_text_from_docx(file_content: bytes) -> str:
    """Extract text from DOCX file"""
    
    if not DOCX_AVAILABLE:
        return "DOCX processing not available. Please install python-docx."
    
    try:
        doc = Document(io.BytesIO(file_content))
        text = ""
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        
        return text
    except Exception as e:
        return f"Error extracting DOCX text: {str(e)}"

def extract_text_from_txt(file_content: bytes) -> str:
    """Extract text from TXT file"""
    try:
        # Try UTF-8 first
        return file_content.decode('utf-8')
    except UnicodeDecodeError:
        try:
            # Fallback to latin-1
            return file_content.decode('latin-1')
        except Exception as e:
            return f"Error reading text file: {str(e)}"

def process_uploaded_file(uploaded_file) -> Dict[str, Any]:
    """Process uploaded file and extract metadata + text"""
    
    file_info = {
        'name': uploaded_file.name,
        'size': uploaded_file.size,
        'type': uploaded_file.type,
        'content': '',
        'word_count': 0,
        'char_count': 0,
        'error': None
    }
    
    try:
        file_content = uploaded_file.read()
        
        # Extract text based on file type
        if uploaded_file.type == "application/pdf":
            text = extract_text_from_pdf(file_content)
        elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            text = extract_text_from_docx(file_content)
        elif uploaded_file.type == "text/plain":
            text = extract_text_from_txt(file_content)
        else:
            text = f"Unsupported file type: {uploaded_file.type}"
            file_info['error'] = "Unsupported file type"
        
        # Clean and analyze text
        cleaned_text = clean_text(text)
        
        file_info.update({
            'content': cleaned_text,
            'word_count': len(cleaned_text.split()),
            'char_count': len(cleaned_text)
        })
        
    except Exception as e:
        file_info['error'] = f"Processing error: {str(e)}"
    
    return file_info

def clean_text(text: str) -> str:
    """Clean and normalize extracted text"""
    if not text:
        return ""
    
    # Remove excessive whitespace
    text = re.sub(r'\s+', ' ', text)
    
    # Remove special characters but keep basic punctuation
    text = re.sub(r'[^\w\s.,!?;:()\[\]{}"\'-]', '', text)
    
    # Normalize line breaks
    text = text.replace('\r\n', '\n').replace('\r', '\n')
    
    return text.strip()

def analyze_text_complexity(text: str) -> Dict[str, Any]:
    """Analyze text complexity metrics"""
    if not text:
        return {}
    
    words = text.split()
    sentences = re.split(r'[.!?]+', text)
    paragraphs = text.split('\n\n')
    
    # Calculate averages
    avg_word_length = sum(len(word) for word in words) / len(words) if words else 0
    avg_sentence_length = len(words) / len(sentences) if sentences else 0
    
    return {
        'word_count': len(words),
        'sentence_count': len(sentences),
        'paragraph_count': len(paragraphs),
        'avg_word_length': round(avg_word_length, 2),
        'avg_sentence_length': round(avg_sentence_length, 2),
        'complexity_score': min(100, int(avg_word_length * 10 + avg_sentence_length))
    }
